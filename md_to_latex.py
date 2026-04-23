#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Assignment Markdown to LaTeX Converter
将作业 Markdown 转换为 LaTeX 并编译为 PDF
"""

import re
import os
import sys
import subprocess
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional

# 尝试导入 pdfplumber（可选，用于图片提取）
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# 尝试导入数学计算库（可选，用于解题辅助）
try:
    import sympy
    HAS_SYMPY = True
except ImportError:
    HAS_SYMPY = False

try:
    import numpy
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False

# 尝试导入图片下载相关库
try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

# 尝试导入 OCR 库（用于图片文字识别）
# 优先使用 EasyOCR（兼容性更好），备选 PaddleOCR
try:
    import easyocr
    HAS_EASYOCR = True
    _easyocr_reader = None
except ImportError:
    HAS_EASYOCR = False
    _easyocr_reader = None

try:
    from paddleocr import PaddleOCR
    HAS_PADDLEOCR = True
    _paddleocr_instance = None
except ImportError:
    HAS_PADDLEOCR = False
    _paddleocr_instance = None

# 尝试导入 python-docx（用于 Word 文档读取）
try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

# 尝试导入 mammoth（用于 Word 转 Markdown，备选方案）
try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

# 默认学生信息（可修改）
DEFAULT_INFO = {
    "name": "姓名",
    "student_id": "学号",
    "class": "班级",
}


def parse_markdown(md_content):
    """解析 Markdown，提取信息"""

    # 提取标题（第一行 # 标题）
    title_match = re.search(r'^#\s+(.+)$', md_content, re.MULTILINE)
    title = title_match.group(1) if title_match else "作业"

    # 提取学生信息
    info = DEFAULT_INFO.copy()
    name_match = re.search(r'\*\*姓名\*\*:\s*(.+)', md_content)
    if name_match:
        info["name"] = name_match.group(1).strip()

    id_match = re.search(r'\*\*学号\*\*:\s*(.+)', md_content)
    if id_match:
        info["student_id"] = id_match.group(1).strip()

    class_match = re.search(r'\*\*班级\*\*:\s*(.+)', md_content)
    if class_match:
        info["class"] = class_match.group(1).strip()

    date_match = re.search(r'\*\*提交日期\*\*:\s*(.+)', md_content)
    submit_date = date_match.group(1).strip() if date_match else datetime.now().strftime("%Y-%m-%d")

    return title, info, submit_date


def extract_images_from_pdf(pdf_path, output_dir="sources"):
    """使用 pdfplumber 从 PDF 中提取图片并保存到 sources/ 目录

    Args:
        pdf_path: PDF 文件路径
        output_dir: 输出目录，默认为 sources/

    Returns:
        list: 提取的图片文件名列表
    """
    if not HAS_PDFPLUMBER:
        print("警告: pdfplumber 未安装，无法提取图片")
        print("安装命令: pip install pdfplumber")
        return []

    images = []
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                if page.images:
                    for img_num, image in enumerate(page.images):
                        # 获取图片边界框
                        x0, top, x1, bottom = image['x0'], image['top'], image['x1'], image['bottom']

                        # 裁剪图片区域
                        cropped = page.crop((x0, top, x1, bottom))

                        # 转换为图片对象
                        image_obj = cropped.to_image()

                        # 生成文件名
                        image_filename = f"page{page_num+1}_img{img_num+1}.png"
                        image_path = output_path / image_filename

                        # 保存图片
                        image_obj.save(str(image_path))
                        images.append(image_filename)
                        print(f"✅ 提取图片: {image_filename}")

    except Exception as e:
        print(f"❌ 提取图片时出错: {e}")

    return images


# ============================================================================
# 图片文字识别功能 (OCR) - 支持从 PNG/JPG/JPEG 提取文字
# ============================================================================

def get_ocr_instance():
    """获取 OCR 实例（懒加载单例模式）
    优先使用 EasyOCR（兼容性更好），备选 PaddleOCR
    """
    global _easyocr_reader, _paddleocr_instance

    # 优先使用 EasyOCR
    if HAS_EASYOCR and _easyocr_reader is None:
        try:
            _easyocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)
            print("✅ 使用 EasyOCR 进行文字识别")
            return _easyocr_reader
        except Exception as e:
            print(f"⚠️ EasyOCR 初始化失败: {e}")

    # 备选：使用 PaddleOCR
    if HAS_PADDLEOCR and _paddleocr_instance is None:
        try:
            _paddleocr_instance = PaddleOCR(
                use_angle_cls=True,
                lang='ch',
                show_log=False,
                use_gpu=False
            )
            print("✅ 使用 PaddleOCR 进行文字识别")
            return _paddleocr_instance
        except Exception as e:
            print(f"⚠️ PaddleOCR 初始化失败: {e}")

    return _easyocr_reader or _paddleocr_instance


def extract_text_from_image(image_path: str) -> str:
    """从图片文件提取文字（支持 PNG, JPG, JPEG）
    优先使用 EasyOCR，备选 PaddleOCR

    Args:
        image_path: 图片文件路径

    Returns:
        str: 识别的文字内容
    """
    ocr = get_ocr_instance()
    if not ocr:
        print("❌ OCR 库未安装或初始化失败")
        print("   安装命令: pip install easyocr")
        return ""

    if not os.path.exists(image_path):
        print(f"❌ 文件不存在: {image_path}")
        return ""

    try:
        # 判断使用哪个 OCR 库
        if HAS_EASYOCR and _easyocr_reader is not None:
            # 使用 EasyOCR
            # 处理中文路径问题：使用 numpy 读取文件，然后用 cv2.imdecode 解码
            import numpy as np
            try:
                import cv2
                # 读取图片文件（支持中文路径）
                with open(image_path, 'rb') as f:
                    img_data = np.frombuffer(f.read(), np.uint8)
                img = cv2.imdecode(img_data, cv2.IMREAD_COLOR)
                if img is None:
                    print(f"❌ 无法读取图片: {image_path}")
                    return ""
                # 使用解码后的图片
                result = _easyocr_reader.readtext(img)
            except ImportError:
                # 如果 cv2 不可用，直接使用路径（可能不支持中文）
                result = _easyocr_reader.readtext(image_path)

            if not result:
                print(f"⚠️ 未从图片中识别到文字: {image_path}")
                return ""
            # 提取文字 (result 是 [(bbox, text, confidence), ...])
            text_lines = [item[1] for item in result]
            return '\n'.join(text_lines)

        elif HAS_PADDLEOCR and _paddleocr_instance is not None:
            # 使用 PaddleOCR
            result = _paddleocr_instance.ocr(image_path, cls=True)
            if not result or not result[0]:
                print(f"⚠️ 未从图片中识别到文字: {image_path}")
                return ""
            # 提取文字
            text_lines = []
            for line in result[0]:
                if line and len(line) >= 2:
                    text_lines.append(line[1][0])
            return '\n'.join(text_lines)

        return ""

    except Exception as e:
        print(f"❌ OCR 识别失败: {e}")
        return ""


def is_image_file(file_path: str) -> bool:
    """检查文件是否为支持的图片格式

    Args:
        file_path: 文件路径

    Returns:
        bool: 是否为图片文件
    """
    image_extensions = ('.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG')
    return file_path.lower().endswith(image_extensions)


def get_file_type(file_path: str) -> str:
    """获取文件类型

    Args:
        file_path: 文件路径

    Returns:
        str: 文件类型 ('pdf', 'image', 'docx', 'doc', 'md', 'unknown')
    """
    file_path_lower = file_path.lower()
    if file_path_lower.endswith('.pdf'):
        return 'pdf'
    elif file_path_lower.endswith(('.png', '.jpg', '.jpeg')):
        return 'image'
    elif file_path_lower.endswith('.docx'):
        return 'docx'
    elif file_path_lower.endswith('.doc'):
        return 'doc'
    elif file_path_lower.endswith('.md'):
        return 'md'
    else:
        return 'unknown'


def extract_content_from_file(file_path: str, output_dir: str = "sources") -> dict:
    """统一接口：从各种文件类型提取内容

    Args:
        file_path: 文件路径（支持 PDF, PNG, JPG, JPEG, DOCX, DOC, MD）
        output_dir: 图片输出目录

    Returns:
        dict: {'text': 文字内容, 'images': 图片列表, 'file_type': 文件类型}
    """
    file_type = get_file_type(file_path)

    result = {
        'text': '',
        'images': [],
        'file_type': file_type
    }

    if file_type == 'pdf':
        # PDF 处理
        print(f"📄 读取 PDF 文件: {file_path}")
        result['images'] = extract_images_from_pdf(file_path, output_dir)

    elif file_type == 'image':
        # 图片 OCR 处理
        print(f"🖼️  读取图片文件: {file_path}")
        print("🔍 正在识别文字（使用 PaddleOCR）...")
        result['text'] = extract_text_from_image(file_path)
        if result['text']:
            print(f"✅ 识别到 {len(result['text'].split())} 个字符")

    elif file_type == 'docx':
        # DOCX 处理
        print(f"📝 读取 DOCX 文件: {file_path}")
        print("🔍 正在提取文字...")
        # 优先使用 mammoth 转换为 Markdown（保留格式）
        result['text'] = extract_docx_to_markdown(file_path)
        if result['text']:
            print(f"✅ 提取到 {len(result['text'].split())} 个字符")

    elif file_type == 'doc':
        # DOC 处理（旧格式）
        print(f"📝 检测到 .doc 文件: {file_path}")
        doc_format = check_doc_format(file_path)

        if doc_format == 'renamed_docx':
            print("ℹ️  检测到这是 .docx 格式（只是扩展名是 .doc）")
            print("   尝试作为 .docx 处理...")
            result['text'] = extract_docx_to_markdown(file_path)
        elif doc_format == 'old_doc':
            print("⚠️  旧版 .doc 格式（Word 97-2003）")
            print("   请在 Microsoft Word 或 LibreOffice 中打开并另存为 .docx 格式")
            print("   或使用 LibreOffice 命令转换:")
            print(f"   soffice --headless --convert-to docx {file_path}")
        else:
            print("⚠️  无法识别的文档格式")
            print("   请将文档转换为 .docx 或 .pdf 格式")

    elif file_type == 'md':
        print(f"📝 Markdown 文件")
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                result['text'] = f.read()
        except Exception as e:
            print(f"❌ 读取失败: {e}")

    else:
        print(f"❌ 不支持的文件类型: {file_path}")

    return result


# ============================================================================
# Word 文档处理功能 (DOCX) - 支持 Word 文档读取
# ============================================================================

def extract_text_from_docx(docx_path: str) -> str:
    """从 .docx 文件提取文字（使用 python-docx）

    Args:
        docx_path: Word 文档路径

    Returns:
        str: 提取的文字内容
    """
    if not HAS_PYTHON_DOCX:
        print("❌ python-docx 未安装")
        print("   安装命令: pip install python-docx")
        return ""

    if not os.path.exists(docx_path):
        print(f"❌ 文件不存在: {docx_path}")
        return ""

    try:
        doc = Document(docx_path)
        text_lines = []

        # 提取段落文字
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_lines.append(paragraph.text)

        # 提取表格文字
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_lines.append(' | '.join(row_text))

        return '\n\n'.join(text_lines)

    except Exception as e:
        print(f"❌ 读取 DOCX 失败: {e}")
        return ""


def extract_docx_to_markdown(docx_path: str) -> str:
    """从 .docx 转换为 Markdown（使用 mammoth，更好保留格式）

    Args:
        docx_path: Word 文档路径

    Returns:
        str: Markdown 内容
    """
    if not HAS_MAMMOTH:
        print("❌ mammoth 未安装，尝试使用 python-docx...")
        print("   安装命令: pip install mammoth")
        # 降级使用基础提取
        return extract_text_from_docx(docx_path)

    if not os.path.exists(docx_path):
        print(f"❌ 文件不存在: {docx_path}")
        return ""

    try:
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file)
            return result.value
    except Exception as e:
        print(f"❌ DOCX 转 Markdown 失败: {e}")
        return ""


def check_doc_format(doc_path: str) -> str:
    """检查 .doc 是否真的是旧格式，还是只是 .docx 重命名

    Args:
        doc_path: .doc 文件路径

    Returns:
        str: 'old_doc', 'renamed_docx', 'unknown'
    """
    try:
        with open(doc_path, 'rb') as f:
            header = f.read(8)
            # ZIP 文件头（.docx 实际上是 ZIP 格式）
            if header.startswith(b'PK\x03\x04'):
                return 'renamed_docx'
            # 旧版 .doc 格式（OLE 格式）
            elif header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                return 'old_doc'
        return 'unknown'
    except Exception:
        return 'unknown'


# ============================================================================
# 图片搜索与下载功能 (Image)
# ============================================================================

def download_image_from_url(url: str, output_path: str) -> bool:
    """从URL下载图片

    Args:
        url: 图片URL
        output_path: 输出文件路径

    Returns:
        bool: 是否下载成功
    """
    if not HAS_REQUESTS:
        print("❌ requests 库未安装，无法下载图片")
        print("安装命令: pip install requests")
        return False

    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        with open(output_path, 'wb') as f:
            f.write(response.content)

        return True
    except Exception as e:
        print(f"❌ 下载失败: {e}")
        return False


def search_images_pexels(query: str, count: int = 5) -> List[Dict[str, str]]:
    """使用 Pexels API 搜索免费高质量图片（需要 API key）

    获取免费 API key: https://www.pexels.com/api/

    Args:
        query: 搜索关键词
        count: 返回结果数量

    Returns:
        List[Dict]: 图片信息列表，每项包含 {url, title, width, height}
    """
    if not HAS_REQUESTS:
        print("❌ 缺少必要库，请安装: pip install requests")
        return []

    # 尝试从配置文件读取 API key
    api_key = None
    config_path = Path.home() / '.claude' / 'assignment_config.json'
    if config_path.exists():
        try:
            import json
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                api_key = config.get('pexels_api_key')
        except Exception:
            pass

    if not api_key:
        print("❌ 未找到 Pexels API key")
        print("   1. 访问 https://www.pexels.com/api/ 获取免费 API key")
        print("   2. 在 ~/.claude/assignment_config.json 中添加:")
        print('      {"pexels_api_key": "YOUR_API_KEY"}')
        return []

    images = []
    try:
        # Pexels API endpoint
        search_url = f"https://api.pexels.com/v1/search?query={query}&per_page={count}&locale=zh-CN"

        headers = {
            'Authorization': api_key
        }

        response = requests.get(search_url, headers=headers, timeout=15)
        response.raise_for_status()

        data = response.json()

        for photo in data.get('photos', [])[:count]:
            # 使用中等尺寸图片
            img_url = photo['src']['medium']
            images.append({
                'url': img_url,
                'title': photo.get('alt', f'Image {len(images)+1}'),
                'width': photo['width'],
                'height': photo['height']
            })

    except Exception as e:
        print(f"⚠️ Pexels API 搜索失败: {e}")

    return images


def search_images_unsplash(query: str, count: int = 5) -> List[Dict[str, str]]:
    """使用 Unsplash API 搜索免费高质量图片（需要 API key）

    获取免费 API key: https://unsplash.com/developers

    Args:
        query: 搜索关键词
        count: 返回结果数量

    Returns:
        List[Dict]: 图片信息列表，每项包含 {url, title, width, height}
    """
    if not HAS_REQUESTS:
        print("❌ 缺少必要库，请安装: pip install requests")
        return []

    # 尝试从配置文件读取 API key
    api_key = None
    config_path = Path.home() / '.claude' / 'assignment_config.json'
    if config_path.exists():
        try:
            import json
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                api_key = config.get('unsplash_access_key')
        except Exception:
            pass

    if not api_key:
        print("❌ 未找到 Unsplash API key")
        print("   1. 访问 https://unsplash.com/developers 获取免费 API key")
        print("   2. 在 ~/.claude/assignment_config.json 中添加:")
        print('      {"unsplash_access_key": "YOUR_ACCESS_KEY"}')
        return []

    images = []
    try:
        # Unsplash API endpoint
        search_url = f"https://api.unsplash.com/search/photos?query={query}&per_page={count}&lang=zh"

        headers = {
            'Authorization': f'Client-ID {api_key}'
        }

        response = requests.get(search_url, headers=headers, timeout=15)
        response.raise_for_status()

        data = response.json()

        for result in data.get('results', [])[:count]:
            # 使用 regular 尺寸图片
            img_url = result['urls']['regular']
            images.append({
                'url': img_url,
                'title': result.get('description') or result.get('alt_description') or f'Image {len(images)+1}',
                'width': result['width'],
                'height': result['height']
            })

    except Exception as e:
        print(f"⚠️ Unsplash API 搜索失败: {e}")

    return images


def search_images_bing(query: str, count: int = 5) -> List[Dict[str, str]]:
    """使用 Bing 搜索图片（备用方案，不推荐）

    注意：此方法可能因防盗链和反爬虫机制而失败。

    Args:
        query: 搜索关键词
        count: 返回结果数量

    Returns:
        List[Dict]: 图片信息列表，每项包含 {url, title, width, height}
    """
    if not HAS_REQUESTS or not HAS_BS4:
        print("❌ 缺少必要库，请安装: pip install requests beautifulsoup4")
        return []

    images = []
    try:
        # Bing Images 搜索 URL
        search_url = f"https://www.bing.com/images/search?q={query}&form=HDRSC2&first=1"

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(search_url, headers=headers, timeout=15)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')

        # 查找图片元素（简化版，实际可能需要调整选择器）
        img_divs = soup.find_all('div', class_='imgpt')[:count]

        for i, div in enumerate(img_divs):
            try:
                img = div.find('img')
                if img and img.get('src'):
                    images.append({
                        'url': img['src'],
                        'title': img.get('alt', f'Image {i+1}'),
                        'width': 'unknown',
                        'height': 'unknown'
                    })
            except Exception:
                continue

    except Exception as e:
        print(f"⚠️ Bing 搜索失败: {e}")

    return images


def search_and_download_images(
    query: str = None,
    url: str = None,
    count: int = 3,
    output_dir: str = "sources",
    auto_select: bool = False
) -> List[str]:
    """搜索并下载图片

    Args:
        query: 搜索关键词（与 url 二选一）
        url: 直接图片URL（与 query 二选一）
        count: 下载数量（仅搜索模式有效）
        output_dir: 输出目录
        auto_select: 是否自动选择前N张图片

    Returns:
        List[str]: 下载的图片文件名列表
    """
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)

    downloaded_files = []
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if url:
        # 模式1：直接从URL下载
        filename = Path(url).name or f"image_{timestamp}.jpg"
        # 如果没有扩展名，默认用 .jpg
        if not Path(filename).suffix:
            filename += ".jpg"

        file_path = output_path / filename

        print(f"📥 正在下载: {url}")
        if download_image_from_url(url, str(file_path)):
            downloaded_files.append(filename)
            print(f"✅ 已保存: {filename}")

    elif query:
        # 模式2：搜索图片
        print(f"🔍 正在搜索: {query}")

        # 优先尝试 Pexels API，然后 Unsplash API，最后 Bing
        images = search_images_pexels(query, count=count)
        if not images:
            images = search_images_unsplash(query, count=count)
        if not images:
            print("⚠️ API 搜索失败，尝试 Bing 搜索（可能不稳定）")
            images = search_images_bing(query, count=count*2)  # 多搜一些备用

        if not images:
            print("❌ 未找到相关图片")
            return []

        # 显示搜索结果
        print(f"📦 找到 {len(images)} 张相关图片:")
        for i, img in enumerate(images[:count*2], 1):
            print(f"   {i}. {img.get('title', 'N/A')} ({img.get('width', '?')}x{img.get('height', '?')})")

        # 选择要下载的图片
        if auto_select:
            selected_indices = list(range(min(count, len(images))))
        else:
            try:
                user_input = input(f"\n选择要下载的图片编号 (1-{min(count, len(images))}, 用逗号分隔): ")
                selected_indices = [int(x.strip()) - 1 for x in user_input.split(',')]
            except (ValueError, EOFError):
                selected_indices = list(range(min(count, len(images))))

        # 下载选中的图片
        for i in selected_indices:
            if 0 <= i < len(images):
                img_url = images[i]['url']
                # 确定文件扩展名
                ext = '.jpg'
                if img_url.endswith('.png'):
                    ext = '.png'
                elif img_url.endswith('.webp'):
                    ext = '.webp'
                elif img_url.endswith('.gif'):
                    ext = '.gif'

                filename = f"image_{len(downloaded_files)+1}_{timestamp}{ext}"
                file_path = output_path / filename

                print(f"📥 下载中 [{i+1}]...")
                if download_image_from_url(img_url, str(file_path)):
                    downloaded_files.append(filename)
                    print(f"✅ 已保存: {filename}")

    return downloaded_files


def insert_image_to_documents(
    md_path: str,
    image_files: List[str],
    position: str = "end",
    section_marker: str = None,
    caption: str = ""
) -> bool:
    """将图片引用插入到 .md 和 .tex 文件

    Args:
        md_path: Markdown 文件路径
        image_files: 图片文件名列表
        position: 插入位置 ("end", "after_section")
        section_marker: 章节标记（position="after_section" 时使用）
        caption: 图片说明

    Returns:
        bool: 是否成功
    """
    if not image_files:
        return False

    tex_path = str(Path(md_path).with_suffix('.tex'))

    # 读取文件
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
    except FileNotFoundError:
        print(f"❌ 文件不存在: {md_path}")
        return False

    # 生成图片引用代码
    md_image_refs = []
    tex_image_refs = []

    for img in image_files:
        md_image_refs.append(f"\n**{caption or '图片'}**：\n")
        md_image_refs.append(f"![{caption or '图片'}](sources/{img})\n")

        tex_image_refs.append("\n\\begin{center}\n")
        tex_image_refs.append(f"\\includegraphics[width=0.7\\textwidth]{{{img}}}\\\\\n")
        tex_image_refs.append(f"\\small{{{caption or '图片'}}}\n")
        tex_image_refs.append("\\end{center}\n")

    # 更新 .md 文件
    if position == "end":
        md_content += ''.join(md_image_refs)
    elif position == "after_section" and section_marker:
        # 在指定章节后插入
        marker_pos = md_content.find(section_marker)
        if marker_pos != -1:
            # 找到章节后的下一个换行
            insert_pos = md_content.find('\n', marker_pos)
            if insert_pos != -1:
                md_content = md_content[:insert_pos+1] + ''.join(md_image_refs) + md_content[insert_pos+1:]

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    # 更新 .tex 文件（如果存在）
    if os.path.exists(tex_path):
        with open(tex_path, 'r', encoding='utf-8') as f:
            tex_content = f.read()

        # 在 \end{document} 前插入
        end_doc_pos = tex_content.rfind('\\end{document}')
        if end_doc_pos != -1:
            tex_content = tex_content[:end_doc_pos] + ''.join(tex_image_refs) + tex_content[end_doc_pos:]

        with open(tex_path, 'w', encoding='utf-8') as f:
            f.write(tex_content)

    return True


def md_to_latex(md_content, title, info, submit_date):
    """将 Markdown 转换为 LaTeX"""

    # 移除开头的标题和学生信息行（会在 LaTeX 中重新生成）
    lines = md_content.split('\n')
    content_start = 0
    for i, line in enumerate(lines):
        if line.startswith('---'):
            content_start = i + 1
            break

    md_body = '\n'.join(lines[content_start:])

    # 转换 Markdown 语法到 LaTeX
    latex_body = convert_markdown_to_latex(md_body)

    # 生成完整 LaTeX 文档
    latex_template = r"""\documentclass[12pt,a4paper]{ctexart}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{geometry}
\usepackage{graphicx}
\usepackage{float}
\usepackage{enumitem}
\geometry{margin=2.5cm}

\title{""" + title + r"""}
\author{""" + info["name"] + r" \quad " + info["student_id"] + r" \quad " + info["class"] + r"""}

\date{""" + submit_date + r"""}

\begin{document}

\maketitle

""" + latex_body + r"""

\end{document}"""

    return latex_template


def convert_markdown_to_latex(md_text):
    """转换 Markdown 语法到 LaTeX"""

    lines = md_text.split('\n')
    latex_lines = []
    in_list = False

    for line in lines:
        # 跳过空行
        if not line.strip():
            if in_list:
                latex_lines.append(r'\end{itemize}')
                in_list = False
            latex_lines.append('')
            continue

        # 标题转换
        if line.startswith('### '):
            latex_lines.append(r'\subsubsection{' + line[4:].strip() + '}')
        elif line.startswith('## '):
            latex_lines.append(r'\subsection{' + line[3:].strip() + '}')
        elif line.startswith('# '):
            latex_lines.append(r'\section{' + line[2:].strip() + '}')
        # 列表转换
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            if not in_list:
                latex_lines.append(r'\begin{itemize}')
                in_list = True
            latex_lines.append(r'\item ' + line.strip()[2:])
        # 加粗转换
        else:
            converted = line
            converted = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', converted)
            # 数学公式转换
            converted = re.sub(r'\$\$(.+?)\$\$', r'[\1]', converted, flags=re.DOTALL)
            converted = re.sub(r'\$(.+?)\$', r'(\1)', converted)
            latex_lines.append(converted)

    if in_list:
        latex_lines.append(r'\end{itemize}')

    return '\n'.join(latex_lines)


def compile_pdf(tex_path, max_runs=2):
    """使用 pdflatex 编译 PDF"""
    try:
        for i in range(max_runs):
            result = subprocess.run(
                ['pdflatex', '-interaction=nonstopmode', tex_path],
                capture_output=True,
                text=True,
                cwd=os.path.dirname(tex_path)
            )
            if result.returncode != 0:
                print(f"编译错误 (第 {i+1} 次):")
                print(result.stderr)
                return False

        # 清理辅助文件
        base_path = str(Path(tex_path).with_suffix(''))
        for ext in ['.aux', '.log', '.out']:
            aux_file = base_path + ext
            if os.path.exists(aux_file):
                os.remove(aux_file)

        return True

    except FileNotFoundError:
        print("错误: 未找到 pdflatex，请安装 MikTeX 或 TeX Live")
        return False


# ============================================================================
# LaTeX 检查与修复功能 (Revise)
# ============================================================================

class LatexIssue:
    """LaTeX 问题类"""
    def __init__(self, line_num: int, issue_type: str, message: str, severity: str = "warning"):
        self.line_num = line_num
        self.issue_type = issue_type
        self.message = message
        self.severity = severity  # "error", "warning", "info"

    def __str__(self):
        severity_symbol = {"error": "❌", "warning": "⚠️", "info": "ℹ️"}
        return f"{severity_symbol.get(self.severity, '•')} 行 {self.line_num}: {self.message} [{self.issue_type}]"


def check_latex_file(tex_path: str) -> List[LatexIssue]:
    """检查 .tex 文件的语法问题

    Args:
        tex_path: LaTeX 文件路径

    Returns:
        List[LatexIssue]: 发现的问题列表
    """
    issues = []

    if not os.path.exists(tex_path):
        issues.append(LatexIssue(0, "file", "文件不存在", "error"))
        return issues

    with open(tex_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 检查特殊字符转义
    special_chars = {
        '%': '百分号',
        '&': '安培符',
        '#': '井号',
        '_': '下划线',
    }

    for line_num, line in enumerate(lines, 1):
        # 跳过注释行
        if line.strip().startswith('%'):
            continue

        for char, name in special_chars.items():
            # 检查未转义的特殊字符（排除已转义的 \\% 和在 \begin、\end 等命令中的）
            if char == '%':
                # 检查未转义的 %（排除 \% 和注释）
                pattern = r'(?<!\\)%'
            elif char == '&':
                # 检查未转义的 &（排除 \&）
                pattern = r'(?<!\\)&'
            elif char == '#':
                # 检查未转义的 #（排除 \#）
                pattern = r'(?<!\\)#'
            elif char == '_':
                # 检查未转义的 _（排除 \_ 和数学模式中的 _）
                # 在数学模式中 _ 是下标，不算错误
                in_math = False
                for i, c in enumerate(line):
                    if c == '$':
                        in_math = not in_math
                    elif c == '_' and not in_math:
                        # 检查是否已转义
                        if i == 0 or line[i-1] != '\\':
                            issues.append(LatexIssue(line_num, "escape", f"未转义的{name}: '{char}' 在非数学模式中", "warning"))
                continue  # _ 单独处理

            matches = re.finditer(pattern, line)
            for match in matches:
                # 排除注释中的 %
                if char == '%' and '\\' in line[:match.start()]:
                    continue
                issues.append(LatexIssue(line_num, "escape", f"未转义的{name}: '{char}'", "warning"))

    # 检查公式配对
    dollar_count = 0
    bracket_count = 0
    for line_num, line in enumerate(lines, 1):
        # 跳过注释行
        if line.strip().startswith('%'):
            continue

        # 检查 $ 配对
        dollar_count += line.count('$')

        # 检查 \[ \] 配对
        bracket_count += line.count(r'\[')
        bracket_count -= line.count(r'\]')

    if dollar_count % 2 != 0:
        issues.append(LatexIssue(0, "math", "行内公式 $ 不配对", "error"))

    if bracket_count != 0:
        issues.append(LatexIssue(0, "math", "显示公式 \\[ \\] 不配对", "error"))

    # 检查环境配对
    environments = {}
    for line_num, line in enumerate(lines, 1):
        # 检查 \begin{...}
        for match in re.finditer(r'\\begin\{([^}]+)\}', line):
            env_name = match.group(1)
            if env_name not in environments:
                environments[env_name] = []
            environments[env_name].append(('begin', line_num))

        # 检查 \end{...}
        for match in re.finditer(r'\\end\{([^}]+)\}', line):
            env_name = match.group(1)
            if env_name not in environments or not environments[env_name]:
                issues.append(LatexIssue(line_num, "environment", f"多余的 \\end{{{env_name}}}", "error"))
            else:
                environments[env_name].pop()

    # 检查未关闭的环境
    for env_name, stack in environments.items():
        if stack:
            for _, line_num in stack:
                issues.append(LatexIssue(line_num, "environment", f"\\begin{{{env_name}}} 未关闭", "error"))

    # 检查图片路径
    for line_num, line in enumerate(lines, 1):
        for match in re.finditer(r'\\includegraphics[^}]*\{([^}]+)\}', line):
            img_path = match.group(1)
            # 解析相对路径
            tex_dir = os.path.dirname(tex_path)
            full_path = os.path.join(tex_dir, img_path)
            if not os.path.exists(full_path):
                issues.append(LatexIssue(line_num, "image", f"图片文件不存在: {img_path}", "error"))

    return issues


def auto_fix_latex(tex_path: str, issues: List[LatexIssue]) -> bool:
    """自动修复常见 LaTeX 问题

    Args:
        tex_path: LaTeX 文件路径
        issues: 问题列表

    Returns:
        bool: 是否成功修复
    """
    if not os.path.exists(tex_path):
        print(f"❌ 文件不存在: {tex_path}")
        return False

    with open(tex_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    fixed_lines = lines.copy()
    offset = 0  # 行号偏移（因为修复可能改变行数）

    for issue in issues:
        if issue.issue_type == "escape" and issue.severity == "warning":
            line_idx = issue.line_num - 1 + offset
            if 0 <= line_idx < len(fixed_lines):
                line = fixed_lines[line_idx]

                # 转义特殊字符
                if '%' in issue.message and '\\%' not in line:
                    fixed_lines[line_idx] = line.replace('%', r'\%')
                elif '&' in issue.message and '\\&' not in line:
                    fixed_lines[line_idx] = line.replace('&', r'\&')
                elif '#' in issue.message and '\\#' not in line:
                    fixed_lines[line_idx] = line.replace('#', r'\#')
                elif '_' in issue.message:
                    # 在非数学模式下转义下划线
                    # 简单处理：在非 $ ... $ 之间的 _ 替换为 \_
                    pass  # 下划线处理较复杂，跳过自动修复

    # 保存修复后的文件
    backup_path = tex_path + '.backup'
    with open(backup_path, 'w', encoding='utf-8') as f:
        f.writelines(lines)
    print(f"📦 备份已保存: {backup_path}")

    with open(tex_path, 'w', encoding='utf-8') as f:
        f.writelines(fixed_lines)

    print(f"✅ 已自动修复 {len(issues)} 个问题")
    return True


def validate_latex_compilation(tex_path: str) -> Tuple[bool, str]:
    """验证 LaTeX 能否成功编译

    Args:
        tex_path: LaTeX 文件路径

    Returns:
        Tuple[bool, str]: (是否成功, 错误信息)
    """
    try:
        result = subprocess.run(
            ['pdflatex', '-interaction=nonstopmode', '-draftmode', tex_path],
            capture_output=True,
            text=True,
            cwd=os.path.dirname(tex_path),
            timeout=30
        )

        # 清理 draftmode 生成的文件
        base_path = str(Path(tex_path).with_suffix(''))
        for ext in ['.aux', '.log']:
            aux_file = base_path + ext
            if os.path.exists(aux_file):
                os.remove(aux_file)

        if result.returncode == 0:
            return True, ""
        else:
            # 提取关键错误信息
            error_output = result.stderr or result.stdout
            # 简化错误输出
            errors = re.findall(r'^!.*$', error_output, re.MULTILINE)
            if errors:
                return False, "\n".join(errors[:5])  # 只返回前5个错误
            return False, "编译失败，请检查日志"

    except FileNotFoundError:
        return False, "未找到 pdflatex，请安装 MikTeX 或 TeX Live"
    except subprocess.TimeoutExpired:
        return False, "编译超时"
    except Exception as e:
        return False, f"验证出错: {str(e)}"


# ============================================================================
# 数学计算辅助功能
# ============================================================================

def math_calculate(expression: str) -> Optional[str]:
    """使用数学库计算表达式

    Args:
        expression: 数学表达式，如 "x**2 - 4 = 0" 或 "diff(x**3, x)"

    Returns:
        Optional[str]: 计算结果字符串
    """
    if HAS_SYMPY:
        try:
            # 尝试解方程
            if '=' in expression:
                # 解方程: x**2 - 4 = 0
                left, right = expression.split('=')
                x = sympy.symbols('x')
                eq = sympy.Eq(sympy.sympify(left), sympy.sympify(right))
                result = sympy.solve(eq, x)
                return f"解方程 {expression} = 0:\n{x} = {result}"
            else:
                # 计算表达式
                result = sympy.sympify(expression)
                return f"{expression} = {result}"
        except Exception as e:
            return f"计算错误: {str(e)}"
    elif HAS_NUMPY:
        try:
            # 使用 numpy 进行数值计算
            result = eval(expression, {"__builtins__": {}}, {"np": numpy, "numpy": numpy})
            return f"{expression} = {result}"
        except Exception as e:
            return f"计算错误: {str(e)}"
    else:
        return "数学计算库未安装。请安装: pip install sympy numpy"


def show_math_help():
    """显示数学计算帮助"""
    help_text = """
📐 数学计算辅助功能

可用库:
"""
    if HAS_SYMPY:
        help_text += "  ✅ SymPy (符号计算)\n"
    else:
        help_text += "  ❌ SymPy (未安装，pip install sympy)\n"

    if HAS_NUMPY:
        help_text += "  ✅ NumPy (数值计算)\n"
    else:
        help_text += "  ❌ NumPy (未安装，pip install numpy)\n"

    help_text += """
使用示例:
  - 解方程: x**2 - 4 = 0  →  x = [2, -2]
  - 微积分: diff(x**3, x)  →  3*x**2
  - 积分: integrate(x**2, x)  →  x**3/3
  - 数值计算: sqrt(16)  →  4.0

Python 代码示例:
  from sympy import symbols, solve, diff, integrate
  x = symbols('x')
  solve(x**2 - 4, x)  # 解方程
  diff(x**3, x)       # 求导
"""
    return help_text


# ============================================================================
# 主函数
# ============================================================================

def main():
    """主函数"""
    # 设置 Windows 控制台编码
    if sys.platform == 'win32':
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

    if len(sys.argv) < 2:
        print("用法: python md_to_latex.py <文件> [选项]")
        print("支持的文件类型:")
        print("  .md                 Markdown 文件（转为 LaTeX 并编译 PDF）")
        print("  .pdf                PDF 文件（使用 --extract 提取内容）")
        print("  .png/.jpg/.jpeg     图片文件（使用 --extract 进行 OCR 识别）")
        print("  .docx               Word 文档（使用 --extract 提取文字）")
        print("  .doc                旧版 Word 文档（需先转换为 .docx）")
        print()
        print("选项:")
        print("  --extract            从文件提取内容（PDF/图片/DOCX）")
        print("  --check              检查 .tex 文件语法")
        print("  --fix                自动修复常见问题")
        print("  --math <expr>        计算数学表达式")
        print("  --math-help          显示数学计算帮助")
        print("  --image <query>      搜索并下载图片（关键词）")
        print("  --image-url <url>    直接从URL下载图片")
        print("  --image-count <n>    图片数量（默认3）")
        print("  --image-pos <pos>    插入位置: end/after_section（默认end）")
        print()
        print("示例:")
        print("  python md_to_latex.py assignment.md")
        print("  python md_to_latex.py homework.pdf --extract")
        print("  python md_to_latex.py worksheet.png --extract")
        print("  python md_to_latex.py report.docx --extract")
        sys.exit(1)

    # 处理命令行选项
    if '--extract' in sys.argv:
        # 从文件提取内容（PDF 或图片）
        file_path = sys.argv[1]
        result = extract_content_from_file(file_path)

        print(f"\n{'='*50}")
        print(f"文件类型: {result['file_type'].upper()}")
        print(f"{'='*50}")

        if result['text']:
            print(f"\n📝 提取的文字内容:\n")
            print(result['text'])
            print(f"\n{'='*50}")
            print(f"✅ 共识别到 {len(result['text'])} 个字符")

        if result['images']:
            print(f"\n🖼️  提取的图片 ({len(result['images'])} 张):")
            for img in result['images']:
                print(f"   - {img}")

        sys.exit(0)

    if '--check' in sys.argv:
        # 检查 LaTeX 语法
        tex_path = sys.argv[1]
        print(f"🔍 检查 LaTeX 文件: {tex_path}")
        issues = check_latex_file(tex_path)

        if not issues:
            print("✅ 没有发现语法问题")
        else:
            print(f"⚠️ 发现 {len(issues)} 个问题:")
            for issue in issues:
                print(f"  {issue}")

        sys.exit(0)

    elif '--fix' in sys.argv:
        # 自动修复 LaTeX 问题
        tex_path = sys.argv[1]
        print(f"🔧 修复 LaTeX 文件: {tex_path}")

        issues = check_latex_file(tex_path)
        fixable_issues = [i for i in issues if i.issue_type == "escape" and i.severity == "warning"]

        if not issues:
            print("✅ 没有发现需要修复的问题")
        elif not fixable_issues:
            print("⚠️ 发现问题但无法自动修复:")
            for issue in issues:
                print(f"  {issue}")
        else:
            print(f"🔧 准备修复 {len(fixable_issues)} 个问题...")
            if auto_fix_latex(tex_path, fixable_issues):
                print("✅ 修复完成")

        sys.exit(0)

    elif '--math' in sys.argv:
        # 数学计算
        idx = sys.argv.index('--math')
        if idx + 1 < len(sys.argv):
            expression = sys.argv[idx + 1]
            print(f"📐 计算: {expression}")
            result = math_calculate(expression)
            print(result)
        else:
            print(show_math_help())
        sys.exit(0)

    elif '--math-help' in sys.argv:
        print(show_math_help())
        sys.exit(0)

    elif '--image' in sys.argv or '--image-url' in sys.argv:
        # 图片搜索与下载
        idx = sys.argv.index('--image') if '--image' in sys.argv else -1
        url_idx = sys.argv.index('--image-url') if '--image-url' in sys.argv else -1

        # 获取参数
        query = None
        url = None
        count = 3
        image_pos = "end"

        if idx != -1 and idx + 1 < len(sys.argv):
            query = sys.argv[idx + 1]
        elif url_idx != -1 and url_idx + 1 < len(sys.argv):
            url = sys.argv[url_idx + 1]

        # 获取可选参数
        if '--image-count' in sys.argv:
            count_idx = sys.argv.index('--image-count')
            if count_idx + 1 < len(sys.argv):
                try:
                    count = int(sys.argv[count_idx + 1])
                except ValueError:
                    pass

        if '--image-pos' in sys.argv:
            pos_idx = sys.argv.index('--image-pos')
            if pos_idx + 1 < len(sys.argv):
                image_pos = sys.argv[pos_idx + 1]

        # 确定输出目录和 markdown 文件
        output_dir = "sources"
        md_path = None

        # 查找当前目录的 .md 文件
        for f in os.listdir('.'):
            if f.endswith('.md'):
                md_path = f
                break

        if not md_path:
            print("❌ 未找到 .md 文件，请指定")
            sys.exit(1)

        # 搜索并下载图片
        downloaded = search_and_download_images(
            query=query,
            url=url,
            count=count,
            output_dir=output_dir
        )

        if downloaded:
            print(f"\n✅ 共下载 {len(downloaded)} 张图片")

            # 询问是否插入到文档
            try:
                insert = input("\n是否插入到文档? (y/n): ").strip().lower()
                if insert == 'y':
                    caption = input("图片说明 (留空则使用默认): ").strip()
                    # 清理可能的代理字符
                    caption = caption.encode('utf-8', errors='ignore').decode('utf-8')
                    insert_image_to_documents(md_path, downloaded, position=image_pos, caption=caption)
                    print("✅ 图片引用已添加到文档")
            except (EOFError, UnicodeEncodeError):
                # 非交互模式或编码错误，自动插入
                insert_image_to_documents(md_path, downloaded, position=image_pos, caption="")
                print("✅ 图片引用已添加到文档")

        sys.exit(0)

    # 默认：Markdown 转 LaTeX 并编译 PDF
    md_path = sys.argv[1]
    if not os.path.exists(md_path):
        print(f"错误: 文件不存在: {md_path}")
        sys.exit(1)

    # 读取 Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 解析并转换
    title, info, submit_date = parse_markdown(md_content)
    latex_content = md_to_latex(md_content, title, info, submit_date)

    # 保存 LaTeX 文件
    base_path = str(Path(md_path).with_suffix(''))
    tex_path = base_path + '.tex'

    with open(tex_path, 'w', encoding='utf-8') as f:
        f.write(latex_content)

    print(f"✅ LaTeX 文件已生成: {tex_path}")

    # 编译 PDF
    print("🔄 正在编译 PDF...")
    if compile_pdf(tex_path):
        pdf_path = base_path + '.pdf'
        print(f"✅ PDF 已生成: {pdf_path}")
    else:
        print("❌ PDF 编译失败")


if __name__ == '__main__':
    main()
